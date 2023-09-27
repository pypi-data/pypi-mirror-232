#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""standard python imports"""

import dataclasses
import json
import os
import re
import zipfile
from datetime import date, datetime
from io import StringIO
from tempfile import gettempdir
from typing import Any, Optional
from urllib.parse import urljoin

import click
import requests
from dateutil.relativedelta import relativedelta
from docx.table import Table
from lxml import etree
from ssp import SSP

from regscale.core.app.api import Api
from regscale.core.app.application import Application
from regscale.core.app.logz import create_logger
from regscale.core.app.public.fedramp.import_fedramp_r4_ssp import (
    parse_and_load_xml_rev4,
)
from regscale.core.app.utils.app_utils import (
    capitalize_words,
    download_file,
    error_and_exit,
    get_current_datetime,
)
from regscale.core.app.utils.regscale_utils import get_user
from regscale.models.app_models.click import regscale_id, regscale_module
from regscale.models.regscale_models import Component, ControlImplementation
from regscale.models.regscale_models.files import File
from regscale.models.regscale_models.interconnects import Interconnects
from regscale.models.regscale_models.leveraged_authorizations import (
    LeveragedAuthorizations,
)
from regscale.models.regscale_models.ports_protocols import PortsProtocols
from regscale.models.regscale_models.requirements import Requirement
from regscale.models.regscale_models.securityplans import SecurityPlan
from regscale.models.regscale_models.system_roles import SystemRoles

logger = create_logger()
namespaces = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


@click.group()
def fedramp():
    """[BETA] Performs bulk processing of FedRAMP files (Upload trusted data only)."""


# FedRAMP Docx Support
# TODO: Improve FedRAMP profile selection.  Prone to mistyping atm.
# TODO: Add stakeholders/teams to created security plan or existing plan.
# TODO: Exporter Takes controls maps to fedramp package.
# TODO: Add extra controls from FedRAMP - HIGH for example. as a prompt
@fedramp.command(context_settings={"show_default": True})
@click.option(
    "--file_name",
    type=click.Path(exists=True, dir_okay=False, file_okay=True),
    required=True,
    prompt="Enter the full file path of the FedRAMP document to ingest to RegScale.",
    help="RegScale will process and load the FedRAMP document.",
)
@click.option(
    "--base_fedramp_profile",
    type=click.STRING,
    required=False,
    help="Enter the name of the RegScale FedRAMP profile to use.",
    default="FedRAMP - High",
)
@click.option(
    "--save_data",
    type=click.BOOL,
    default=False,
    required=False,
    help="Whether to save the data as a JSON file.",
)
def load_fedramp_docx(
    file_name: click.Path, base_fedramp_profile: click.STRING, save_data: click.BOOL
):
    """
    [BETA] Convert a FedRAMP docx file to a RegScale SSP.
    """
    process_fedramp_docx(file_name, base_fedramp_profile, save_data)


@fedramp.command()
@click.option(
    "--file_name",
    type=click.Path(exists=True, dir_okay=False, file_okay=True),
    required=True,
    prompt="Enter the file name of the FedRAMP JSON document to process.",
    help="RegScale will process and load the FedRAMP document.",
)
@click.option(
    "--submission_date",
    type=click.DateTime(formats=["%Y-%m-%d"]),
    default=str(date.today()),
    required=True,
    prompt="Enter the submission date of this FedRAMP document.",
    help=f"Submission date, default is today: {date.today()}.",
)
@click.option(
    "--expiration_date",
    type=click.DateTime(formats=["%Y-%m-%d"]),
    default=str((datetime.now() + relativedelta(years=3)).date()),
    required=True,
    prompt="Enter the expiration date of this FedRAMP document.",
    help=f"Expiration date, default is {str((datetime.now() + relativedelta(years=3)).date())}.",
)
def load_fedramp_oscal(file_name, submission_date, expiration_date):
    """
    [BETA] Convert a FedRAMP OSCAL SSP json file to a RegScale SSP.
    """
    if not expiration_date:
        today_dt = date.today()
        expiration_date = date(today_dt.year + 3, today_dt.month, today_dt.day)

    process_fedramp_oscal_ssp(file_name, submission_date, expiration_date)


@fedramp.command()
@click.option(
    "--file-path",
    "-f",
    type=click.Path(exists=True),
    help="File to upload to RegScale.",
    required=True,
)
@click.option(
    "--catalogue_id",
    "-c",
    type=click.INT,
    help="The RegScale ID # of the catalogue to use for controls in the profile.",
    required=True,
)
def import_fedramp_ssp_xml_rev4(file_path: click.Path, catalogue_id: click.INT):
    """
    [BETA] Import FedRAMP Revision 4 SSP XML into RegScale
    """

    logger.info("Importing FedRAMP Revision 4 SSP XML into RegScale")
    parse_and_load_xml_rev4(file_path, catalogue_id)


def decode_access_level(key: str) -> str:
    """
    Decodes the access level from the FedRAMP document
    :param str key: Key used to decode the access level
    :return: Access level as a string
    :rtype: str
    """
    access_levels = {
        "P": "Privileged",
        "NP": "Non-Privileged",
        "NLA": "No Logical Access",
    }
    return access_levels.get(key, "Unknown")


def create_responsible_roles(app: Application, table_data: list, ssp_id: int) -> None:
    """
    [BETA] Inserts the actual the Responsible Roles into the Security Plan.
    :param Application app: Application object
    :param list table_data: list of dicts
    :param int ssp_id: RegScale SSP ID
    :return: None
    """
    roles = [
        table
        for table in table_data
        if "Role" in table.keys() and "Internal or External" in table.keys()
    ]
    logger.info(f"Found {len(roles)} Responsible Roles")
    for role in roles:
        userId = app.config.get("userId")
        accessLevel = decode_access_level(
            role.get(
                "Privileged (P), Non-Privileged (NP), or No Logical Access (NLA)",
                "Unknown",
            )
        )
        systemRole = SystemRoles.get_or_create(
            app=app,
            role_name=role.get("Role"),
            ssp_id=ssp_id,
            roleType=role.get("Internal or External", "Internal"),
            accessLevel=accessLevel,
            sensitivityLevel=role.get("Sensitivity Level", "Not Applicable"),
            assignedUserId=userId,
            privilegeDescription=role.get("Authorized Privileges", "Not Applicable"),
            securityPlanId=ssp_id,
            functions=[role.get("Functions Performed")],
            createdById=userId,
        )


def post_responsible_roles(app: Application, table_data: list, ssp_id: int) -> dict:
    """
    [BETA] Insert the Responsible Roles into the Security Plan.
    :param Application app: Application object
    :param list table_data: list of dicts
    :param int ssp_id: RegScale SSP ID
    :return: dict of the control to role mappings
    :rtype: dict
    """
    data = [
        table for table in table_data if "Control Summary Information" in table.keys()
    ]
    system_roles = list()

    unique_values = set(system_roles)
    ctrl_roles = dict()
    for obj in data:
        control = list(obj.keys())[0] if type(obj) == dict and obj.keys() else None
        for value in obj.values():
            if type(value) is str and value.startswith("Responsible Role:"):
                # Extract the roles from the line
                role = value.split(":", 1)[1].strip()
                # Split the roles by comma and strip whitespace
                if role.lower() not in unique_values:
                    unique_values.add(role.lower())
                    system_roles.append(role)
                userId = app.config.get("userId")
                systemRole = SystemRoles.get_or_create(
                    app=app,
                    role_name=role,
                    ssp_id=ssp_id,
                    roleType="Internal",
                    accessLevel="Privileged",
                    sensitivityLevel="Not Applicable",
                    assignedUserId=userId,
                    privilegeDescription=role,
                    securityPlanId=ssp_id,
                    functions=[role],
                    createdById=userId,
                )
                if control:
                    friendly_control_id = get_friendly_control_id(control)
                    if control in ctrl_roles.keys():
                        ctrl_roles[friendly_control_id].append(systemRole.id)
                    else:
                        ctrl_roles[friendly_control_id] = [systemRole.get("id")]
    return ctrl_roles


def process_fedramp_oscal_ssp(
    file_path: click.Path, submission_date: date, expiration_date: date
) -> None:
    """
    OSCAL FedRAMP to RegScale SSP
    :param click.Path file_path: A click file path object to the oscal file
    :param date submission_date: The Submission date YYYY-MM-DD
    :param date expiration_date: The Expiration date YYYY-MM-DD
    :raises: FileNotFoundError if provided file doesn't exist
    :raises: JSONDecodeError if unable to read the JSON file
    :return: None
    """
    app = Application()
    api = Api(app)
    config = app.config
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            ssp_dict = json.load(file)
    except FileNotFoundError:
        error_and_exit(f"File not found!\n{file_path}")
    except json.JSONDecodeError as jex:
        logger.error(
            "JSONDecodeError, something is wrong with the file: %s\n%s", file_path, jex
        )

    # Create SSP
    create_ssp(api, config, ssp_dict, submission_date, expiration_date)


def check_profile(api: Api, config: dict, title: str) -> list:
    """
    Check if the profile exists in RegScale
    :param Api api: The api instance
    :param dict config: The application configuration
    :param str title: The title of the profile in question
    :raises: ValueError if the provided title doesn't exist in RegScale
    :return: List of filtered profiles
    :rtype: list
    """
    profiles_response = api.get(config["domain"] + "/api/profiles/getList")
    if profiles_response.status_code == 200:
        profiles = profiles_response.json()
    filtered = [dat for dat in profiles if dat["name"] == title]
    if not filtered:
        raise ValueError(
            f"The profile {title} does not exist in RegScale, \
                please create it and re-run this task."
        )
    return filtered[0]["id"]


def create_port(api: Api, config: dict, dat: PortsProtocols) -> None:
    """Create a port and protocol for a component

    :param Api api: An api instance
    :param dict config: Configuration
    :param PortsProtocols dat: Port and protocol data
    :param dict headers: headers for api call
    :return: None
    """

    existing_ports = api.get(
        url=config["domain"]
        + f"/api/portsProtocols/getAllByParent/{dat.parentId}/components",
    ).json()
    if dat not in [PortsProtocols.from_dict(port) for port in existing_ports]:
        # Check if obj exists
        port_res = api.post(
            url=config["domain"] + "/api/portsProtocols",
            data=json.dumps(dat.__dict__),
        )
        if port_res.status_code == 200:
            logger.info("Port and Protocol for component %i added!", dat.parentId)
        else:
            logger.warning(
                "Unable to post Port and Protocol: %s.",
                json.dumps(dat),
            )


def create_ssp_components(
    api: Api, config: dict, components: dict, ssp_id: int
) -> None:
    """
    Creates SSP Components
    :param Api api: The API instance
    :param dict config: The application's configuration
    :param dict components: The components
    :param int ssp_id: The ID of the SSP in RegScale
    :return: None
    """
    component_types = [
        "hardware",
        "software",
        "policy",
        "service",
        "process",
        "procedure",
        "compliance artifact",
    ]
    ports = set()
    for component in components:
        comp_type = (
            component["type"]
            if component["type"].lower() in component_types
            else "compliance artifact"
        )
        status = "Inactive/Retired"
        if component["status"]["state"] == "operational":
            status = "Active"

        comp = Component(
            title=component["title"],
            securityPlansId=ssp_id,
            componentType=comp_type,
            lastUpdatedById=config["userId"],
            createdById=config["userId"],
            cmmcExclusion=False,
            componentOwnerId=config["userId"],
            description=component["description"],
            status=status,
        )

        # save component
        cmp_id = None
        url = urljoin(config["domain"], "/api/components")
        cmp_response = api.post(
            url=url,
            json=comp.dict(),
        )
        if cmp_response.ok:
            cmp = cmp_response.json()
            cmp_id = cmp["id"]
            logger.info(
                "Successfully posted new component# %i as %s for ssp# %i.",
                cmp_id,
                cmp["title"],
                ssp_id,
            )
        if cmp_id and "protocols" in component.keys():
            for protocol in component["protocols"]:
                ports_protocols = PortsProtocols(
                    service="",
                    usedBy="",
                    parentId=cmp_id,
                    purpose=component["type"],
                    startPort=int(protocol["port-ranges"][0]["start"]),
                    endPort=int(protocol["port-ranges"][0]["end"]),
                    protocol=protocol["name"],
                    parentModule="components",
                    lastUpdatedById=config["userId"],
                    createdById=config["userId"],
                )
                ports.add(ports_protocols)

        if component["type"].lower() == "interconnection" and cmp_id:
            # Create ports and protocols object
            ports_protocols = PortsProtocols(
                service="",
                usedBy="",
                parentId=cmp_id,
                purpose=component["type"],
                startPort=0,
                endPort=0,
                protocol="",
                parentModule="components",
                lastUpdatedById=config["userId"],
                createdById=config["userId"],
            )
            ports_protocols.parentId = cmp_id
            ports_protocols.purpose = component["type"]
            # loop through properties to find port number
            if "props" in component.keys():
                for prop in component["props"]:
                    if prop["name"] == "information":
                        ports_protocols.purpose = prop["value"]
                    if prop["name"] == "port":
                        ports_protocols.startPort = int(prop["value"])
                        ports_protocols.endPort = int(prop["value"])
            ports.add(ports_protocols)
        create_component_mapping(api, config, ssp_id, cmp_id)
    if ports:
        for dat in ports:
            create_port(api, config, dat)


def create_component_mapping(api: Api, config: dict, ssp_id: int, cmp_id: int) -> None:
    """
    Create Component Mapping
    :param Api api: The api instance.
    :param dict config: The application configuration.
    :param int ssp_id: The SSP ID.
    :param int cmp_id: The component ID.
    :param dict headers: A dict of headers for requests.
    :return: None
    """
    mapping = {
        "securityPlanId": ssp_id,
        "componentId": cmp_id,
        "isPublic": True,
        "createdById": config["userId"],
        "lastUpdatedById": config["userId"],
    }
    mapping_response = api.post(
        url=config["domain"] + "/api/componentmapping",
        data=json.dumps(mapping),
    )
    if mapping_response.status_code != 200:
        logger.warning("Unable to post Mapping Response: %s.", mapping)


def create_ssp_stakeholders(
    api: Api, config: dict, ssp_id: int, ssp_dict: dict
) -> None:
    """
    Create Stakeholders in RegScale
    :param Api api: The api instance.
    :param dict config: The application configuration.
    :param int ssp_id: The SSP ID.
    :param dict ssp_dict: An SSP Dictionary.
    :return: None
    """
    parties = ssp_dict["system-security-plan"]["metadata"]["parties"]
    filtered_parties = list(filter(lambda x: x["type"] == "person", parties))
    for party in filtered_parties:
        title = [dat["value"] for dat in party["props"] if dat["name"] == "job-title"]
        phone = [dat["number"] for dat in party["telephone-numbers"]]
        email = list(party["email-addresses"])
        addresses = list(party["addresses"]) if "addresses" in party.keys() else None
        stakeholder = {
            "name": party["name"],
            "title": title[0] if title else "",
            "phone": phone[0] if phone else "",
            "email": email[0] if email else "",
            "address": addresses[0]["addr-lines"][0]
            + " "
            + addresses[0]["city"]
            + " "
            + addresses[0]["state"]
            + ", "
            + addresses[0]["postal-code"]
            if addresses
            else "",
            "otherID": party["uuid"],
            "notes": email[0] if email else "",
            "parentId": ssp_id,
            "parentModule": "securityplans",
        }
        post_stakeholder(api, config, stakeholder)


def post_stakeholder(api: Api, config: dict, stakeholder: dict) -> None:
    """Post Stakeholders to RegScale

    :param Api api: API instance
    :param dict config: An application configuration
    :param dict stakeholder: A stakeholder dictionary
    :return: None
    """
    headers = {
        "authorization": config["token"],
        "accept": "application/json",
        "Content-Type": "application/json",
    }
    response = api.post(
        url=config["domain"] + "/api/stakeholders",
        headers=headers,
        data=json.dumps(stakeholder),
    )
    logger.info("Posting Stakeholder: %s", stakeholder)
    if response.status_code != 200:
        logger.warning("Unable to post stakeholder: %s.", stakeholder)


def create_ssp_control_implementations(
    api: Api,
    config: dict,
    ssp_id: int,
    controls: dict,
    ssp_dict: dict,
) -> None:
    """
    Create the control implementations from the oscal SSP object
    :param Api api: The api instance.
    :param dict config: The application configuration.
    :param int ssp_id: The SSP ID.
    :param dict controls: A dict of existing controls in RegScale.
    :param dict ssp_dict: An SSP Dictionary.
    :return: None
    """
    if not controls:
        return
    control_implementations = ssp_dict["system-security-plan"][
        "control-implementation"
    ]["implemented-requirements"]

    for implementation in control_implementations:
        status = "Not Implemented"

        for prop in implementation["props"]:
            if prop["name"] == "implementation-status":
                status = capitalize_words(prop["value"].replace("-", " "))
                if prop["value"].lower() == "implemented":
                    status = "Fully Implemented"
                if prop["value"].lower() == "partial":
                    status = "Partially Implemented"

        control_id = [
            control["controlID"]
            for control in controls
            if control["controlId"].lower() == implementation["control-id"].lower()
        ][0]
        imp = ControlImplementation(
            parentId=ssp_id,
            parentModule="securityplans",
            controlID=control_id,
            controlOwnerId=config["userId"],
            lastUpdatedById=config["userId"],
            createdById=config["userId"],
            status=status,
        )
        # Post Implementation
        response = post_regscale_object(api=api, config=config, obj=imp)
        imp_id = None
        if not response.raise_for_status and response.status_code == 200:
            imp_id = response.json()["id"]
        if imp_id:
            # TODO: Add statements
            pass


def post_regscale_object(
    api: Api, config: dict, obj: Any, endpoint="controlimplementation"
) -> requests.Response:
    """
    Post RegScale control implementation
    :param Api api: API instance
    :param dict config: Application config
    :param Any obj: data object
    :param str endpoint: Endpoint to use in RegScale
    :raises: General error if unable to post object to RegScale
    :return: Response from API call to RegScale
    :rtype: requests.Response
    """
    response = None
    headers = {
        "authorization": config["token"],
        "accept": "application/json",
        "Content-Type": "application/json",
    }
    try:
        response = api.post(
            config["domain"] + f"/api/{endpoint}",
            headers=headers,
            data=json.dumps(obj)
            if isinstance(obj, dict)
            else json.dumps(dataclasses.asdict(obj)),
        )
    except Exception as ex:
        logger.error("Unable to Post %s: %s to RegScale.\n%s", endpoint, obj, ex)

    return response


def create_ssp(
    api: Api, config: dict, ssp_dict: dict, submission_date: date, expiration_date: date
) -> int:
    """
    Create a basic SSP in RegScale
    :param Api api: The api instance.
    :param dict config: The application configuration.
    :param dict ssp_dict: An SSP Dictionary.
    :param date submission_date: The Submission date YYYY-MM-DD
    :param date expiration_date: The Expiration date YYYY-MM-DD
    :return: A newly created RegScale security plan id.
    :rtype: int
    """
    existing_ssps = []
    metadata = ssp_dict["system-security-plan"]["metadata"]
    system = ssp_dict["system-security-plan"]["system-characteristics"]
    fedramp_profile = get_profile(
        ssp_dict["system-security-plan"]["import-profile"]["href"]
    )["profile"]
    profile_id = check_profile(api, config, fedramp_profile["metadata"]["title"])
    components = ssp_dict["system-security-plan"]["system-implementation"]["components"]
    dt_format = "%Y-%m-%d %H:%M:%S"
    ssp_payload = {
        "uuid": ssp_dict["system-security-plan"]["uuid"],
        "systemName": system.get("system-name", None),  # Required
        "planInformationSystemSecurityOfficerId": config["userId"],
        "planAuthorizingOfficialId": config["userId"],
        "systemOwnerId": config["userId"],
        "otherIdentifier": system["system-ids"][0]["id"],
        "confidentiality": capitalize_words(
            system["system-information"]["information-types"][0][
                "confidentiality-impact"
            ]["selected"].split("-")[2]
        ),  # Required
        "integrity": capitalize_words(
            system["system-information"]["information-types"][0]["integrity-impact"][
                "selected"
            ].split("-")[2]
        ),  # Required
        "availability": capitalize_words(
            system["system-information"]["information-types"][0]["availability-impact"][
                "selected"
            ].split("-")[2]
        ),  # Required
        "status": capitalize_words(
            system["status"].get("state", "operational")
        ),  # Required
        "description": system.get("description", None),
        "dateSubmitted": submission_date.strftime(dt_format),
        "approvalDate": (submission_date + relativedelta(years=1)).strftime(
            dt_format
        ),  # User must be changed
        "expirationDate": expiration_date.strftime(dt_format),
        "systemType": "Major Application",  # User must change
        "purpose": metadata.get("", None),
        "conditionsOfApproval": metadata.get("", None),
        "environment": metadata.get("", None),
        "lawsAndRegulations": metadata.get("", None),
        "authorizationBoundary": metadata.get("", None),
        "networkArchitecture": metadata.get("", None),
        "dataFlow": metadata.get("", None),
        "overallCategorization": capitalize_words(
            system["security-sensitivity-level"].split("-")[2]
        ),
        "maturityTier": metadata.get("", None),
        "createdById": config["userId"],
        "hva": False,
        "practiceLevel": metadata.get("", None),
        "processLevel": metadata.get("", None),
        "cmmcLevel": metadata.get("", None),
        "cmmcStatus": metadata.get("", None),
        "isPublic": True,
        "executiveSummary": metadata.get("", None),
        "recommendations": metadata.get("", None),
        "importProfile": metadata.get("version", "fedramp1.1.0-oscal1.0.0"),
        "parentId": profile_id,
        "parentModule": "profiles",
    }
    logger.warning("Unknown System Type, defaulting to %s.", ssp_payload["systemType"])
    logger.warning("Unknown HVA status, defaulting to %r.", ssp_payload["hva"])

    existing_ssp_response = api.get(
        url=urljoin(config["domain"], "/api/securityplans/getList")
    )
    if (
        not existing_ssp_response.raise_for_status()
        and existing_ssp_response.status_code == 200
    ):
        existing_ssps = existing_ssp_response.json()

    if system["system-name"] in {ssp["systemName"] for ssp in existing_ssps}:
        dat = {
            ssp["id"]
            for ssp in existing_ssps
            if ssp["systemName"] == system["system-name"]
        }
        click.confirm(
            f"This SSP Title already exists in the system, \
                SSP: {dat.pop() if len(dat) < 2 else dat}.  Would you still like to continue?",
            abort=True,
        )

    response = api.post(
        url=urljoin(config["domain"], "/api/securityplans"), json=ssp_payload
    )
    if not response.raise_for_status() and response.status_code == 200:
        logger.info("SSP Created with an id of %i!", response.json()["id"])
        ssp_id = response.json()["id"]
    controls = []
    controls_response = api.get(
        config["domain"] + f"/api/profilemapping/getByProfile/{profile_id}"
    )
    if controls_response.status_code == 200:
        controls = controls_response.json()
    create_ssp_components(api, config, components, ssp_id)
    create_ssp_control_implementations(api, config, ssp_id, controls, ssp_dict)
    create_ssp_stakeholders(api, config, ssp_id, ssp_dict)
    # update_ssp_contacts(api, config, ssp_id, ssp_dict)

    return ssp_id


def get_profile(url: str) -> dict:
    """
    Downloads the FedRAMP profile
    :param str url: A profile URL.
    :return: A dictionary with the profile json data.
    :rtype: dict
    """
    dl_path = download_file(url)
    with open(dl_path, encoding="utf-8") as json_file:
        data = json.load(json_file)
    return data


def get_tables(document) -> list:
    """
    Return all document tables
    :param document: document object
    :return: List of all document tables
    :rtype: list
    """
    tables = list(document.tables)
    for t_table in document.tables:
        for row in t_table.rows:
            for cell in row.cells:
                tables.extend(iter(cell.tables))
    return tables


# TODO - replace these dangerous default mutables or document why they are used
def get_xpath_data_detailed(
    tables: Table, key: str, ident: str, xpath: str, count_array: list = [2, 3, 4]
) -> dict:
    """
    Use Xpath to pull data from XML tables.
    :param Table tables: XML tables
    :param str key: specific key in XML table
    :param str ident:
    :param str xpath: xpath of the element
    :param list count_array: array of numbers, default is [2, 3, 4]
    :return: Dictionary of items found
    :rtype: dict
    """
    tables = iter(tables)
    confidentiality = None
    integrity = None
    availability = None
    for t_var in tables:
        if key in t_var._element.xml:
            f = StringIO(t_var._element.xml)
            tree = etree.parse(f)
            tags = tree.xpath(xpath, namespaces=namespaces)
            for p_var in tags:
                t_tags = p_var.xpath("//w:r/w:t", namespaces=namespaces)
                count = 0
                for t_var in t_tags:
                    if t_var.text == ident or count > 0:
                        count += 1
                        if count == count_array[0]:
                            confidentiality = t_var.text
                        if count == count_array[1]:
                            integrity = t_var.text
                        if count == count_array[2]:
                            availability = t_var.text
    return {
        "type": key,
        "nist_ident": ident,
        "confidentiality": confidentiality,
        "integrity": integrity,
        "availability": availability,
    }


def get_contact_info(tables, key: str, xpath: str) -> dict:
    """
    Use Xpath to pull data from XML tables
    :param tables: XML tables
    :param str key: key to look for
    :param str xpath: xpath of the element
    :raises: IndexError if an index error is encountered during iteration
    :return: Dictionary of sorted data
    :rtype: dict
    """
    idents = [
        "Name",
        "Title",
        "Company / Organization",
        "Address",
        "Phone Number",
        "Email Address",
    ]
    dat = {}

    def loop_and_update(element_list) -> dict:
        """
        Loop through the element list and update the data dictionary
        :param element_list:
        :return: Updated dictionary
        :rtype: dict
        """
        value = ""
        field = None
        if idents:
            count = 0
            field = idents.pop(0)

            while count < len(element_list) - 1:
                # if t_tags[new_count].text[len(t_tags[new_count].text)-1] == ' ':
                #     field = ''.join([field, t_tags[new_count+1].text])
                #     new_count = new_count + 1
                if element_list[count].text == field:
                    value = "".join([value, element_list[count + 1].text])
                    dat[field] = value
                    value = ""
                    count += 1
                try:
                    if element_list[count + 1].text in idents:
                        field = idents.pop(0)
                    else:
                        if field in dat:
                            dat[field] = "".join(
                                [dat[field], element_list[count + 1].text]
                            )
                        count += 1
                except IndexError:
                    logger.debug("Unable to continue, index error on row: %i.", count)
                continue

    tables = iter(tables)

    tag_data = []
    for _, t_enum in enumerate(tables):
        if key in t_enum._element.xml:
            f_var = StringIO(t_enum._element.xml)
            tree = etree.parse(f_var)
            tags = tree.xpath(xpath, namespaces=namespaces)
            for tag in tags:
                p_tags = tag.xpath("//w:p", namespaces=namespaces)
                for p_var in p_tags:
                    t_tags = p_var.xpath("//w:r/w:t", namespaces=namespaces)
                    for tags in t_tags:
                        tag_data.append(tags)
    loop_and_update(tag_data)

    return dat


def get_base_contact(ssp, key: Optional[str] = "Point of Contact") -> dict:
    """
    Gets contact information
    :param ssp: SSP file
    :param str key: key to parse for, defaults to 'Point of Contact'
    :return: dictionary with contact information
    :rtype: dict
    """
    result = {}
    name = None
    title = None
    company = None
    address = None
    phone = None
    email = None
    for table in ssp.document.tables:
        for i, row in enumerate(table.rows):
            text = (cell.text.strip() for cell in row.cells)
            if i == 0:
                keys = tuple(text)
                if not title:
                    dat = [x.strip() for x in keys if x].pop().split("\n")
                    title = dat[1] if len(dat) > 1 else dat[0]
                continue
            if key == keys[0]:
                # Use python-docx to get the embedded xml text.
                for i, cell in enumerate(table._cells):
                    if i == 3:
                        name = cell.text
                        result["name"] = name
                    if i == 5:
                        title = cell.text
                        result["title"] = title
                    if i == 7:
                        company = cell.text
                        result["company"] = company
                    if i == 9:
                        address = cell.text
                        result["address"] = address
                    if i == 11:
                        phone = cell.text
                        result["phone"] = phone
                    if i == 13:
                        email = cell.text
                        result["email"] = email
                return result


def post_interconnects(app: Application, table_data: list, regscale_ssp: dict) -> None:
    """
    Interconnects map to SSP in RegScale
    :param Application app: Application object
    :param table_data: List of tables
    :param SecurityPlan regscale_ssp: SecurityPlan object
    :return: None
    """
    # TODO: Update ssp_id to pull the full ssp object down.
    api = Api(app)
    user_id = app.config["userId"]
    user = get_user(api=api, user_id=user_id)
    if len(user) == 0:
        user = {"firstName": user_id, "lastName": "userId: "}
    key = "SP* IP Address and Interface"
    existing_interconnects = []
    dat = [table for table in table_data if key in table.keys()]
    existing_interconnect_response = api.get(
        app.config["domain"]
        + f"/api/interconnections/getAllByParent/{regscale_ssp['id']}/securityplans"
    )
    if not existing_interconnect_response.raise_for_status() and (
        existing_interconnect_response.headers.get("content-type")
        == "application/json; charset=utf-8"
    ):
        existing_interconnects = existing_interconnect_response.json()
    format = "%Y-%m-%d %H:%M:%S"
    for interconnect in dat:
        interconnection = Interconnects(
            name=interconnect[key],
            # connectionType=clean_dict['Data Direction(incoming, outgoing, or both)'],
            ao=f"{user['lastName']}, {user['firstName']}",
            aOId=user_id,
            interconnectOwner="Eaton, Bryan",
            interconnectOwnerId=user_id,
            dateCreated=get_current_datetime(),
            dateLastUpdated=get_current_datetime(),
            lastUpdatedById=user_id,
            createdBy=f"{user['lastName']}, {user['firstName']}",
            createdById=user_id,
            description=interconnect["Information Being Transmitted"]
            if "Information Being Transmitted" in interconnect.keys()
            else "",
            parentId=regscale_ssp["id"],
            parentModule="securityplans",
            agreementDate=get_current_datetime(),
            expirationDate=(datetime.now() + relativedelta(years=3)).strftime(format),
            status="Approved",
            organization=regscale_ssp["systemName"],
            categorization=regscale_ssp["overallCategorization"],
            connectionType="Web Service or API",
            authorizationType="Interconnect Security Agreement (ISA)",
        )
        if interconnection.name + interconnection.description not in {
            inter["name"] + inter["description"] for inter in existing_interconnects
        }:
            post_regscale_object(
                api=api,
                config=app.config,
                obj=interconnection.dict(),
                endpoint="interconnections",
            )


def post_ports(app: Application, table_data: list, ssp_id: int) -> None:
    """
    Ports map to interconnects
    :param Application app: Application object
    :param list table_data: list of tables
    :param int ssp_id: RegScale SSP ID
    :return: None
    """
    api = Api(app)
    key = "Ports (TCP/UDP)*"
    dat = [table for table in table_data if "Protocols" in table.keys()]
    existing_ports = []
    existing_ports_response = api.get(
        app.config["domain"]
        + f"/api/portsProtocols/getAllByParent/{ssp_id}/securityplans"
    )
    if not existing_ports_response.raise_for_status() and (
        existing_ports_response.headers.get("content-type")
        == "application/json; charset=utf-8"
    ):
        existing_ports = existing_ports_response.json()

    for protocol in dat:
        ports_protocols = PortsProtocols(
            service=protocol["Services"],
            usedBy=protocol["Used By"],
            parentId=ssp_id,
            purpose=protocol["Purpose"],
            startPort="".join(c for c in protocol[key] if c.isdigit())
            or "".join(c for c in protocol["Protocols"] if c.isdigit()),
            endPort="".join(c for c in protocol[key] if c.isdigit())
            or "".join(c for c in protocol["Protocols"] if c.isdigit()),
            protocol="".join(c for c in protocol[key] if not c.isdigit())
            or "".join(c for c in protocol["Protocols"] if c.isdigit()),
            parentModule="securityplans",
            lastUpdatedById=app.config["userId"],
            createdById=app.config["userId"],
        )
        ports_protocols.protocol = (
            ports_protocols.protocol.strip().replace("(", "").replace(")", "")
        )
        if ports_protocols not in {
            PortsProtocols.from_dict(port) for port in existing_ports
        }:  # Hashable class
            post_regscale_object(
                api, app.config, ports_protocols.__dict__, endpoint="portsProtocols"
            )


def get_current_implementations(app: Application, regscale_id: int) -> list[dict]:
    """Pull current implementations for a given regscale id
    :param Application app: Application instance
    :param int regscale_id: RegScale ID
    :return: List of dictionaries
    """
    current_imps = []
    api = Api(app)
    try:
        current_imps_response = api.get(
            url=app.config["domain"]
            + f"/api/controlImplementation/getAllByPlan/{regscale_id}",
            params=({"skip_check": True}),
        )
        if not current_imps_response.raise_for_status():
            current_imps = current_imps_response.json()
    except requests.HTTPError:  # This endpoint returns 404 when empty.
        current_imps = []
    return current_imps


def get_friendly_control_id(control_number: str) -> str:
    """Get friendly control id from control number
    :param str control_number: Control number
    :return: Friendly control id
    :rtype: str
    """
    # exp = r"^.*?\([^\d]*(\d+)[^\d]*\).*$"
    # the above regex allows for a denial of service attack
    # the below regex should mitigate that
    exp = r"\((\d+)\)"
    friendly_control_id = control_number.lower()
    if match := re.search(exp, control_number):
        friendly_control_id = f"{control_number[:match.regs[1][0] - 1].strip()}.{match.groups()[0]}".lower()
    return friendly_control_id


# flake8: noqa: C901
def post_implementations(
    app: Application,
    ssp_obj: dict,
    regscale_ssp: dict,
    mapping: dict,
    ctrl_roles: dict,
    save_data: bool = False,
) -> None:
    """
    Post implementations to RegScale
    :param Application app: Application object
    :param dict ssp_obj: SecurityPlan object (python-docx)
    :param dict regscale_ssp: RegScale ssp
    :param dict mapping: mapping
    :param dict ctrl_roles: Control roles
    :param bool save_data: Whether to save data to a file
    :raises ValueError: If there are more controls in the source document than in the base profile
    :return: None
    """
    api = Api(app)
    implementation_status = "Not Implemented"
    implementation_text = ""
    responsibility = None
    new_implementations = []
    friendly_control_id = None
    has_requirements = False
    load_additional_controls = False
    exp = r"^.*?\([^\d]*(\d+)[^\d]*\).*$"
    logger.info(
        "Attempting to post %i controls from this FedRAMP SSP Document to RegScale!",
        len(ssp_obj.control_list),
    )
    current_imps = get_current_implementations(app, regscale_ssp["id"])
    if len(ssp_obj.control_list) > len(mapping):
        # TODO: Handle this case from base catalog
        raise ValueError(
            "There are more controls in the source document \
                            than in the base profile!"
        )

    if len(ssp_obj.control_list) < len(mapping) > len(current_imps):
        profile_name = api.get(
            urljoin(app.config["domain"], f"/api/profiles/{mapping[0]['profileID']}")
        ).json()["name"]
        load_additional_controls = click.prompt(
            f"You have less FedRAMP document controls than the current profile, would you like to add all controls associated with {profile_name}? (y,n)",
            type=click.BOOL,
        )

    if len(current_imps) > 0:
        logger.info(
            "This RegScale Security plan already has %i implementations..",
            len(current_imps),
        )
    mapped_controls_log = []
    unmapped_controls_log = []
    implemented_controls_log = []
    for fedramp_control in ssp_obj.control_list:
        if fedramp_control.control_origination:
            # Shared = Shared
            # Customer = Configured/Provided Customer
            # Provider = Corporate/System Specific
            # Not Applicable
            # TODO:
            if "Shared".lower() in fedramp_control.control_origination[0].lower():
                responsibility = "Shared"
            elif "Customer".lower() in fedramp_control.control_origination[0].lower():
                responsibility = "Customer"
            elif "Provider".lower() in fedramp_control.control_origination[0].lower():
                responsibility = "Provider"
            else:
                responsibility = fedramp_control.control_origination[0]
        if (
            fedramp_control.implementation_status
            and fedramp_control.implementation_status[0]
            in [
                "Alternative Implementation",
                "Implemented",
            ]
        ):
            implementation_status = "Fully Implemented"
        elif "Partially Implemented" in fedramp_control.implementation_status:
            implementation_status = "Partially Implemented"
        else:
            implementation_status = (
                fedramp_control.implementation_status[0]
                if fedramp_control.implementation_status
                else "Not Implemented"
            )
        friendly_control_id = get_friendly_control_id(fedramp_control.number)

        if control := [
            control
            for control in mapping
            if control["controlId"].lower() == friendly_control_id
        ]:
            controlID = control[0]["controlID"]
        else:
            logger.warning(
                'Unable to map control "%s" to RegScale', friendly_control_id.upper()
            )
            unmapped_controls_log.append(friendly_control_id.upper())
            continue
        if len(fedramp_control.parts) > 1:
            implementation_text = "<br>".join(
                fedramp_control.part(x).text for x in fedramp_control.parts
            )
        else:
            try:
                implementation_text = fedramp_control.part(None).text
            except IndexError:
                implementation_text = ""
        imp_response = None
        implementation = ControlImplementation(
            parentId=regscale_ssp["id"],
            parentModule="securityplans",
            controlOwnerId=app.config["userId"],
            status=implementation_status,
            controlID=controlID,
            responsibility=responsibility,
            implementation=implementation_text,
            systemRoleId=ctrl_roles.get(friendly_control_id)[0]
            if isinstance(ctrl_roles, dict) and friendly_control_id in ctrl_roles.keys()
            else None,
        )
        if "Req" in fedramp_control.number:
            # generate child control
            parent_security_control_id = [
                control["controlID"]
                for control in mapping
                if control["controlId"] == friendly_control_id.split()[0]
            ][0]
            # update current implementation list
            current_imps = get_current_implementations(
                app=app, regscale_id=regscale_ssp["id"]
            )
            parent_security_control = [
                imp
                for imp in current_imps
                if imp["controlID"] == parent_security_control_id
            ][0]
            has_requirements = True
        if has_requirements:
            # generate n requirements for new control
            for part in fedramp_control.parts:
                implementation_text = fedramp_control.part(part).text
                title = f"{friendly_control_id.split()[0]} - Req. {part}"
                requirement = Requirement(
                    id=0,
                    description=implementation_text.split("\n")[0],
                    implementation=implementation_text,
                    title=title,
                    lastUpdatedById=app.config["userId"],
                    status=implementation_status,
                    controlID=parent_security_control_id,
                    parentId=parent_security_control["id"],
                    parentModule="controls",
                    requirementOwnerId=app.config["userId"],
                    createdById=app.config["userId"],
                )
                existing_requirement = api.get(
                    url=app.config["domain"]
                    + f"/api/requirements/getByParent/{parent_security_control['id']}/controls"
                ).json()
                if title not in {req["title"] for req in existing_requirement}:
                    logger.info("Posting Requirement: %s", title)
                    post_regscale_object(
                        api=api,
                        config=app.config,
                        obj=requirement,
                        endpoint="requirements",
                    )
                else:
                    logger.info("Requirement %s already exists, skipping...", title)
            has_requirements = False
            implementation_text = None
        else:
            filter = controlID not in {
                imp["controlID"] for imp in new_implementations
            } and controlID not in {imp["controlID"] for imp in current_imps}

            if filter:  # No duplicate implementations
                mapped_controls_log.append(friendly_control_id.upper())
                logger.info("Posting implementation: %s.", fedramp_control.number)
                imp_response = post_regscale_object(
                    api=api, config=app.config, obj=implementation.dict()
                )
                if imp_response.status_code == 200:
                    new_implementations.append(imp_response.json())
            else:
                logger.warning(
                    "Implementation already exists in this RegScale Security Plan: %s, skipping...",
                    fedramp_control.number,
                )
                implemented_controls_log.append(friendly_control_id.upper())
        #  update parameters
        if imp_response:
            for parameter in fedramp_control.parameters:
                # TODO: Bug in 4.12.2 api call. A terrible workaround \
                # using controlParameters endpoint
                param = {
                    "id": 0,
                    "controlImplementationId": imp_response.json()["id"],
                    "name": parameter.split(":")[0],
                    "dataType": "String",
                    "value": parameter.split(":")[1],
                    "createdById": app.config["userId"],
                    "lastUpdatedById": app.config["userId"],
                    "externalPropertyName": None,
                    "oscalNamespaceMapping": [],
                }

                post_regscale_object(api, app.config, param, "parameters")

    # TODO: Add python-ssp license blurb on https://regscale.readme.io/docs/overview
    if load_additional_controls:
        load_non_matched_profile_controls(
            app, regscale_ssp=regscale_ssp, mapping=mapping
        )
    if save_data:
        with open("control_implementation.log", "w") as f:
            f.write("|*** Unmapped Controls ***|\n")
            f.write("------------------------------\n")
            f.write("\n".join(unmapped_controls_log))
            f.write("\n")
            f.write("------------------------------\n")
            f.write("|*** Mapped Controls ***|\n")
            f.write("------------------------------\n")
            f.write("\n".join(mapped_controls_log))
            f.write("\n")
            f.write("------------------------------\n")
            f.write("|*** Already Implemented Controls ***|\n")
            f.write("------------------------------\n")
            f.write("\n".join(implemented_controls_log))
            f.write("\n")
            f.write("------------------------------\n")


def load_non_matched_profile_controls(
    app: Application, regscale_ssp: dict, mapping: dict
) -> None:
    """Load controls from a given profile mapping that are not matched by the document

    :param Application app: Application instance
    :param dict regscale_ssp: RegScale SSP as a dictionary
    :param dict mapping: Profile mapping
    :rtype: None
    """
    api = Api(app)
    current_imps = get_current_implementations(app, regscale_ssp["id"])
    if ssp := [
        ssp
        for ssp in api.get(
            url=urljoin(app.config["domain"], "/api/securityplans/getList")
        ).json()
        if ssp["title"] == regscale_ssp["systemName"]
    ]:
        ssp_id = ssp[0]["id"]
        existing_controls = {imp["controlID"] for imp in current_imps}
        controls_to_add = [
            control
            for control in mapping
            if control["controlID"] not in existing_controls
        ]
        logger.info("Adding %i additional controls from profile", len(controls_to_add))
        existing_control_ids = {imp["controlID"] for imp in current_imps}
        for control in controls_to_add:
            if control["controlID"] not in existing_control_ids:
                implementation = ControlImplementation(
                    parentId=ssp_id,
                    parentModule="securityplans",
                    controlOwnerId=app.config["userId"],
                    status="Not Implemented",
                    controlID=control["controlID"],
                    responsibility=None,
                    implementation=None,
                ).__dict__
                logger.info("Posting implementation: %s.", control["controlId"])
                post_regscale_object(api, app.config, implementation)


def post_attachments(api: Api, link: str, ssp_obj: SSP, regscale_ssp: dict) -> None:
    """
    Download and post Attachments to RegScale
    :param Api api: API object
    :param str link: link to download file
    :param SSP ssp_obj: SSP object
    :param dict regscale_ssp: dictionary of RegScale SSP
    :raises: General error if unable to download file
    :return: None
    """
    try:
        dl_path = download_file(link["link"])
        logger.info("Posting linked image to RegScale.. %s", link)
        File.upload_file_to_regscale(
            file_name=(dl_path.absolute()),
            parent_id=regscale_ssp["id"],
            parent_module="securityplans",
            api=api,
        )

    except Exception as ex:
        logger.warning("Unable to download file: %s\n%s", link, ex)


def posted_embedded_attachments(api: Api, ssp_obj: SSP, regscale_ssp: dict) -> None:
    """
    Find and post embedded picture files to RegScale
    :param Api api: API object
    :param SSP ssp_obj: SSP object
    :param dict regscale_ssp: RegScale SSP
    :return None:
    """
    filename = ssp_obj.source
    with zipfile.ZipFile(filename, mode="r") as archive:
        file_dump_path = gettempdir() + os.sep + "imagedump"
        for file in archive.filelist:
            if (
                file.filename.startswith("word/media/") and file.file_size > 200000
            ):  # 200KB filter
                archive.extract(file, path=file_dump_path)
        # Create directories in case they do not exist.
        media_path = file_dump_path + os.sep + "word" + os.sep + "media"
        if not os.path.exists(media_path):
            os.makedirs(media_path)
        for filename in os.listdir(file_dump_path + os.sep + "word" + os.sep + "media"):
            full_file_path = os.path.join(
                file_dump_path + os.sep + "word" + os.sep + "media", filename
            )
            if os.path.isfile(full_file_path):
                logger.info("Posting embedded image to RegScale... %s", full_file_path)
                File.upload_file_to_regscale(
                    file_name=full_file_path,
                    parent_id=regscale_ssp["id"],
                    parent_module="securityplans",
                    api=api,
                )


def post_links(
    config: dict, api: Api, ssp_obj: SSP, regscale_ssp: dict, post_embeds: bool = True
) -> None:
    """
    Use Xpath to pull data from XML tables'
    :param dict config: Application config
    :param Api api: Api object
    :param SSP ssp_obj: SSP object
    :param dict regscale_ssp: RegScale SSP
    :param bool post_embeds: Whether to post embedded items to RegScale
    :return: None
    """
    # Find and post attachments
    # TODO: Add issue for link titles to pull from XML table instead of hyperlink name.
    # TODO: Add table title to link information
    attachments = []
    titles = []
    if post_embeds:
        posted_embedded_attachments(api, ssp_obj, regscale_ssp)
    for table in ssp_obj.document.tables:
        if table._cells and "Identification Number" in table.cell(0, 0).text.strip():
            previous = None
            id_text = None
            title = None
            link_text = None
            link_date = None
            dat = {}
            for index, e_id in enumerate(
                table._element.xpath(".//w:r/w:t")
            ):  # Loop through every column + 1 record
                dat_text = e_id.text.strip()
                previous = "" if index == 0 else previous
                if dat_text:
                    if previous.lower() == "link":
                        id_text = dat_text
                        dat["id"] = id_text
                        link_text = None
                    elif id_text:
                        title = dat_text
                        id_text = None
                    elif validate_date_str(dat_text):
                        title = " ".join([title, previous]) if title else previous
                        dat["title"] = title.strip()
                        link_date = e_id.text.strip()
                        dat["date"] = link_date.strip()
                        title = None
                    elif "date" in dat:
                        link_text = e_id.text.strip()
                        dat["link"] = link_text
                        titles.append(dat)
                        previous = "link"
                        link_text = None
                    previous = e_id.text.strip() if len(dat) != 4 else "link"
                    if len(dat) == 4:
                        dat = {}
                    link_date = None
            titles.reverse()
            for link in table._element.xpath(".//w:hyperlink"):
                inner_run = link.xpath("w:r", namespaces=link.nsmap)[0]

                # matches = [tit['title'] for tit in titles if tit['link'] == inner_run.text]
                # if matches:\
                # TODO: Improve this.
                title = titles.pop()["title"] if titles else inner_run.text
                # print link relationship id
                r_id = link.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                )
                # print link URL
                attachments.append(
                    {"title": title, "link": ssp_obj.document._part.rels[r_id]._target}
                )
    # Post Links
    existing_links = api.get(
        config["domain"]
        + f"/api/links/getAllByParent/{regscale_ssp['id']}/securityplans"
    ).json()
    for reg_link in list({v["link"]: v for v in attachments}.values()):
        dat_text = {
            "id": 0,
            "url": reg_link["link"],
            "title": reg_link["title"],
            "parentID": regscale_ssp["id"],
            "parentModule": "securityplans",
        }
        if reg_link["link"] not in {link["url"] for link in existing_links}:
            post_regscale_object(api, config, dat_text, endpoint="links")
            post_attachments(
                api=api, link=reg_link, ssp_obj=ssp_obj, regscale_ssp=regscale_ssp
            )
        else:
            logger.warning(
                "%s already exists in Security Plan, skipping..", reg_link["link"]
            )

    # Post Files


def validate_date_str(date_text, fmt="%m/%d/%Y") -> bool:
    """
    Validate provided text is a date in mm/dd/yyyy format
    :param str date_text: Date as a string
    :param str fmt: date format of the date_text, defaults to %m/%d/%Y
    :return: Whether provided text can be converted as a date
    :rtype: bool
    """
    try:
        datetime.strptime(date_text, fmt)
    except ValueError:
        return False
    return True


def get_text(document, full_text, start_header, start_text, end_text) -> str:
    """
    Parses text from a document
    :param document: Document to parse
    :param full_text: Full text of the document
    :param start_header: Starting header
    :param start_text: Where the text starts
    :param end_text: Where the text ends
    :return: String parsed from document
    :rtype: str
    """
    description_num = [
        i for i, j in enumerate(full_text) if j.lower() == start_header.lower()
    ].pop()
    description_text = []
    keep_going = True
    start_run = False
    while keep_going:
        for index, para in enumerate(document.paragraphs):
            if index > description_num:
                for run in para.runs:
                    if run.text.lower().strip() == start_text.lower().strip():
                        start_run = True
                        break
                    if start_run and run.text.strip() != end_text.strip():
                        description_text.append(run.text)
                    if run.text.lower().strip() == end_text.strip().lower():
                        start_run = False
                        keep_going = False
                        break
    return "".join(description_text)


def gather_stakeholders(app: Application, tables: list, regscale_ssp: dict, ssp):
    """Gather Stakeholders

    :param Application app: An application instance
    :param list tables: A list of tables from the XML document.
    :param dict regscale_ssp: A dict of RegScale SSP data.
    :param ssp: A dict of docx SSP data.
    """
    api = Api(app)
    pocs = []
    # Add Management POC
    management_poc = _extracted_from_gather_stakeholders_(
        tables,
        "Owner Information",
        pocs,
        "Information System Management Point of Contact",
    )
    if management_poc:
        pocs.append(management_poc)
    # Add Information POC
    information = get_contact_info(
        tables,
        key="Information System Technical Point of Contact",
        xpath="//w:tbl/w:tr",
    )
    pocs.extend([information, get_base_contact(ssp)])
    # Add CSP POC
    csp_poc = _extracted_from_gather_stakeholders_(
        tables,
        "AO Point of Contact",
        pocs,
        "CSP Name Internal ISSO (or Equivalent) Point of Contact",
    )
    existing_stakeholders = []
    pocs.append(csp_poc)
    existing_stakeholders_response = api.get(
        url=app.config["domain"]
        + f"/api/stakeholders/getAllByParent/{regscale_ssp['id']}/securityplans"
    )
    if existing_stakeholders_response.status_code == 200:
        existing_stakeholders = existing_stakeholders_response.json()
    pocs_inserted = []
    # Filter non-dict objects
    pocs = [poc for poc in pocs if isinstance(poc, dict)]
    for poc in pocs:
        poc = {k.lower(): v for k, v in poc.items()}  # Make this case-insensitive.
        email = name = title = phone = ""
        if "email" in (keys := [key.lower() for key in poc.keys()]):
            email = poc[
                (_ := _check_if_string_in_list_of_string("email", keys)).lower()
            ]
        if "name" in (keys := [key.lower() for key in poc.keys()]):
            name = poc[(_ := _check_if_string_in_list_of_string("name", keys)).lower()]
        if "title" in (keys := [key.lower() for key in poc.keys()]):
            name = poc[(_ := _check_if_string_in_list_of_string("title", keys)).lower()]
        if "phone" in (keys := [key.lower() for key in poc.keys()]):
            name = poc[(_ := _check_if_string_in_list_of_string("phone", keys)).lower()]

        stakeholder = {
            "name": name,
            "title": title,
            "phone": phone,
            "email": email,
            "address": poc.get("address", "") if "address" in poc.keys() else "",
            "parentId": regscale_ssp["id"],
            "parentModule": "securityplans",
        }
        if "name" in poc.keys():
            if poc["name"].strip() not in pocs_inserted and poc["name"].strip() not in {
                guy["name"] for guy in existing_stakeholders if "name" in guy.keys()
            }:
                post_stakeholder(api=api, config=app.config, stakeholder=stakeholder)
                pocs_inserted.append(poc["name"].strip())


def _check_if_string_in_list_of_string(string: str, list_of_strings: list) -> str:
    """
    Check if a string is in a list of strings
    :param str string: The string to check
    :return: the found string
    :rtype: str
    """
    dat = any(string in s for s in list_of_strings)
    if dat:
        return {s for s in list_of_strings if string in s}.pop()
    return ""


def _extracted_from_gather_stakeholders_(
    tables, key: str, pocs: list, key2: str
) -> dict:
    """
    Extract system owner and add to pocs
    :param tables:
    :param str key: Key used to find system owner
    :param list pocs: List of points of contacts from a SSP
    :param str key2: Second key to find system owner details
    :return: Dictionary of system owner details
    :rtype: dict
    """
    system_owner = get_contact_info(tables, key=key, xpath="//w:tbl/w:tr")
    pocs.append(system_owner)
    return get_contact_info(tables, key=key2, xpath="//w:tbl/w:tr")


def post_leveraged_authorizations(
    app: Application, table_data: list, ssp_id: int
) -> None:
    """
    Function to post leveraged authorizations
    :param Application app: Application instance
    :param list table_data: Data used to post to RegScale
    :param int ssp_id: RegScale SSP ID #
    :return: None
    """
    # LeveragedAuthorizations
    KEY = "Leveraged Information System Name"
    data = [table for table in table_data if KEY in table.keys()]
    for la in data:
        if not la.get("Leveraged Information System Name", None):
            continue
        leveraged_auth = LeveragedAuthorizations(
            title=la.get(KEY, " "),
            servicesUsed=la.get("Leveraged Service Provider Owner"),
            dateAuthorized=la.get("Date Granted"),
            createdById=app.config.get("userId"),
            securityPlanId=ssp_id,
            ownerId=app.config.get("userId"),
            lastUpdatedById=app.config.get("userId"),
            dateLastUpdated=get_current_datetime(),
            dateCreated=get_current_datetime(),
        )
        try:
            response_json = LeveragedAuthorizations.insert_leveraged_authorizations(
                app, leveraged_auth
            )
            logger.info(
                f"Leveraged Authorizations for {la.get('Leveraged Information System Name')} created in RegScale."
            )
        except Exception as e:
            logger.error(e)


def process_fedramp_docx(
    fedramp_file_path: click.Path, base_fedramp_profile: str, save_data: bool = False
) -> None:
    """
    Convert a FedRAMP file to a RegScale SSP
    :param click.Path fedramp_file_path: The click file path object
    :param str base_fedramp_profile: base fedramp profile
    :raises: ValueError if unable to get profile from RegScale
    :raises: IndexError, AttributeError if unable to get profile mapping by name in RegScale
    :raises: ValueError if unable to get profile mapping RegScale
    :return: None
    """
    # If list of controls is more than profile mapping, make sure i get them from somewhere? Get base catalog from profile.
    app = Application()
    api = Api(app)
    ssp = SSP(fedramp_file_path)
    document = ssp.document
    full_text = []
    [full_text.append(para.text) for para in document.paragraphs]
    description = get_text(
        document,
        full_text,
        start_header="Information System Components and Boundaries",
        start_text="System Description:",
        end_text="Types of Users",
    )
    environment = get_text(
        document,
        full_text,
        start_header="SYSTEM ENVIRONMENT AND INVENTORY",
        start_text="PRODUCTION ENVIRONMENT: (IMPLEMENTED)",
        end_text="Data Flow",
    )
    count = 0
    title = None
    system_status = "Other"
    system_type = "Major Application"
    confidentiality = "Low"
    integrity = "Low"
    availability = "Low"
    tables = get_tables(document)
    # information_types = get_information_types(tables)
    security_objective = get_xpath_data_detailed(
        tables,
        key="Security Objective",
        ident="Confidentiality",
        xpath="//w:tbl/w:tr",
        count_array=[2, 4, 6],
    )

    # TODO: add cloud deployment models

    availability = security_objective["availability"].split(" ")[0]
    confidentiality = security_objective["confidentiality"].split(" ")[0]
    integrity = security_objective["integrity"].split(" ")[0]
    table_data = []
    for table in ssp.document.tables:
        for i, row in enumerate(table.rows):
            checked = False
            rem = row._element
            check_boxes = rem.xpath(".//w14:checked")
            text = (cell.text.strip() for cell in row.cells)
            if check_boxes:
                for checks in check_boxes:
                    if checks.items()[0][1] == "1":
                        count = count + 1
                        checked = True
            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                if not title:
                    dat = [x.strip() for x in keys if x].pop().split("\n")
                    title = dat[1] if len(dat) > 1 else dat[0]
                continue
            row_data = dict(zip(keys, text))
            if checked:
                if "System Status" in row_data:
                    system_status = row_data["System Status"]
                if "Service Provider Architecture Layers" in row_data:
                    system_type = row_data["Service Provider Architecture Layers"]
            row_data["checked"] = checked
            row_data["element"] = rem
            table_data.append(row_data)
    status = "Operational" if "in production" in system_status else "Other"
    # Links are posted to links mapped to ssp
    # post_links(app, table_data, ssp_id)
    # Parts will go in implementation fields.
    profile_name = base_fedramp_profile
    logger.info("Using the %s profile to import controls.", profile_name)
    profile_response = api.get(
        url=urljoin(app.config["domain"], "/api/profiles/getList")
    )

    if profile_response.ok:
        profiles = profile_response.json()
        logger.info(profiles)
    else:
        profiles = []
        logger.error("Unable to get profiles from RegScale.")
        profile_response.raise_for_status()
    try:
        profile = None
        for profile_obj in profiles:
            if profile_obj["name"] == profile_name:
                profile = profile_obj
        if profile is None:
            raise ValueError(f"Unable to find profile: {profile_name}")
        profile_mapping = api.get(
            urljoin(
                app.config["domain"],
                f"/api/profileMapping/getByProfile/{profile['id']}",
            )
        ).json()
    except (IndexError, AttributeError) as ex:
        logger.error("Unable to continue, %s is not found!\n%s", profile_name, ex)
    if len(profile_mapping) == 0:
        raise ValueError(
            f"Unable to continue, please load {profile_name} with controls!",
            profile_name,
        )
    today_dt = date.today()
    expiration_date = date(today_dt.year + 3, today_dt.month, today_dt.day)
    dt_format = "%Y-%m-%d %H:%M:%S"
    text = []
    regscale_ssp = SecurityPlan(
        dateSubmitted=get_current_datetime(),
        expirationDate=expiration_date.strftime(dt_format),
        approvalDate=expiration_date.strftime(dt_format),
        parentId=profile["id"],
        parentModule="profiles",
        systemName=title or "Unable to determine System Name",
        confidentiality=capitalize_words(confidentiality) or "Moderate",
        integrity=capitalize_words(integrity) or "Moderate",
        availability=capitalize_words(availability) or "Moderate",
        status=status or "Operational",
        createdById=app.config["userId"],
        lastUpdatedById=app.config["userId"],
        systemOwnerId=app.config["userId"],
        planAuthorizingOfficialId=app.config["userId"],
        planInformationSystemSecurityOfficerId=app.config["userId"],
        systemType=system_type or "Major Application",
        overallCategorization=ssp.system_security_level or "Moderate",
        description=description or "Unable to determine System Description",
        environment=environment or "",
        executiveSummary=f"Revision: {ssp.revision}",
    ).dict()
    if regscale_ssp.get("status") != "Operational":
        regscale_ssp.explanationForNonOperational = (
            "Unable to determine status from SSP during FedRAMP .docx import."
        )
    existing_security_plans_reponse = api.get(
        url=urljoin(app.config["domain"], "/api/securityplans/getList"),
    )
    existing_security_plans = []
    if not existing_security_plans_reponse.ok:
        logger.info(f"No Security Plans found")
    else:
        existing_security_plans = existing_security_plans_reponse.json()

    if regscale_ssp["systemName"].lower() not in {
        sys["title"].lower() for sys in existing_security_plans
    }:
        regscale_ssp_response = api.post(
            url=urljoin(app.config["domain"], "/api/securityplans"),
            json=regscale_ssp,
        )
        if regscale_ssp_response.ok:
            regscale_ssp = regscale_ssp_response.json()
        else:
            regscale_ssp_response.raise_for_status()
            logger.error(f"Unable to create Security Plan: {regscale_ssp}")
    if "id" in regscale_ssp and not regscale_ssp["id"]:
        plan_list = api.get(
            urljoin(app.config["domain"], "/api/securityplans/getList")
        ).json()
        if regscale_ssp["systemName"] in [reg["title"] for reg in plan_list]:
            regscale_ssp["id"] = [
                sec["id"]
                for sec in plan_list
                if sec["title"] == regscale_ssp["systemName"]
            ][0]
    create_responsible_roles(app, table_data, ssp_id=regscale_ssp["id"])
    ctrl_roles = post_responsible_roles(app, table_data, ssp_id=regscale_ssp["id"])
    gather_stakeholders(app, tables, regscale_ssp, ssp)
    post_interconnects(app, table_data, regscale_ssp)
    post_ports(app, table_data, ssp_id=regscale_ssp["id"])
    post_links(config=app.config, api=api, ssp_obj=ssp, regscale_ssp=regscale_ssp)
    post_implementations(
        app=app,
        ssp_obj=ssp,
        regscale_ssp=regscale_ssp,
        mapping=profile_mapping,
        ctrl_roles=ctrl_roles,
        save_data=save_data,
    )
    post_leveraged_authorizations(app, table_data, ssp_id=regscale_ssp["id"])
