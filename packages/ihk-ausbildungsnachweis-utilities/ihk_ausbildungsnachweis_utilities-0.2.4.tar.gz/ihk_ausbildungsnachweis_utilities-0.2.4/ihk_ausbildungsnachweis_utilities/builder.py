# Copyright (C) 2023 twyleg
import re
import shutil
import subprocess
import lxml.etree as ET
import tempfile

from typing import Dict
from linecache import getline
from pathlib import Path
from docx import Document
from xmlschema import XMLSchema, XMLSchemaValidationError

FILE_DIR = Path(__file__).parent

LIBREOFFICE_BINARY = "libreoffice"

DEFAULT_TEMPLATE_FILE_NAME = "template_ausbildungsnachweis.docx"


class InvalidInputFileContentException(Exception):
    def __init__(self, msg: str, line_number: int, line: str):
        super().__init__(f'Invalid input: "{msg}", Line {line_number}: {line}')


def __validate_input_for_build_pdf(
    ausbildungsnachweis_input_xml_filepath: Path, ausbildungsachweis_output_pdf_dirpath: Path, template_filepath: Path
) -> None:
    if not ausbildungsnachweis_input_xml_filepath.is_file():
        raise FileNotFoundError(ausbildungsnachweis_input_xml_filepath)

    if not ausbildungsachweis_output_pdf_dirpath.is_dir():
        raise FileNotFoundError(ausbildungsachweis_output_pdf_dirpath)

    if not template_filepath.is_file():
        raise FileNotFoundError(template_filepath)


def __get_template_file_path(template_path: Path) -> Path:
    if template_path.is_file():
        return template_path
    elif template_path.is_dir():
        return Path(template_path / DEFAULT_TEMPLATE_FILE_NAME)
    else:
        raise FileNotFoundError(template_path)


def build_ausbildungsnachweis_pdf(
    ausbildungsnachweis_input_xml_filepath: Path,
    ausbildungsachweis_output_pdf_dirpath: Path,
    template_filepath: Path = FILE_DIR / "templates/template_ausbildungsnachweis.docx",
) -> Path:
    __validate_input_for_build_pdf(
        ausbildungsnachweis_input_xml_filepath, ausbildungsachweis_output_pdf_dirpath, template_filepath
    )

    tmp_dirpath = Path(tempfile.mkdtemp())
    tmp_docx_filepath = tmp_dirpath / f"{ausbildungsnachweis_input_xml_filepath.stem}.docx"
    ausbildungsachweis_output_pdf_filepath = ausbildungsachweis_output_pdf_dirpath / f"{tmp_docx_filepath.stem}.pdf"

    try:
        __fill_template(ausbildungsnachweis_input_xml_filepath, tmp_docx_filepath, template_filepath)
    except XMLSchemaValidationError as e:
        reason = e.reason if e.reason else "Unknown reason!"
        raise InvalidInputFileContentException(
            reason, e.sourceline, getline(str(ausbildungsnachweis_input_xml_filepath), e.sourceline)
        )

    __convert_docx_to_pdf(tmp_docx_filepath, ausbildungsachweis_output_pdf_dirpath)

    return ausbildungsachweis_output_pdf_filepath


def __read_node(root_node: ET._Element, child_node_name: str, normalized=True) -> str:
    child_node = root_node.find(child_node_name)
    if child_node is None:
        raise RuntimeError(f'Expected child node "{child_node_name}" is missing.')
    node_text_raw = child_node.text if child_node.text else ""
    if normalized:
        p = re.compile(r"^\s+")
        m = p.match(node_text_raw)
        if m:
            node_text_raw = re.sub(m.group(), "\n", node_text_raw)
    return node_text_raw


def __read_input(input_filepath: Path) -> Dict[str, str]:
    schema = XMLSchema(FILE_DIR / "schemas/ausbildungsnachweis.xsd")

    element_tree = ET.parse(input_filepath)
    schema.validate(element_tree)  # type: ignore
    root_node = element_tree.getroot()

    input: Dict[str, str] = {
        "vorname": __read_node(root_node, "vorname"),
        "name": __read_node(root_node, "name"),
        "nr": __read_node(root_node, "nr"),
        "ausbildungsjahr": __read_node(root_node, "ausbildungsjahr"),
        "ausbildungswoche_von": __read_node(root_node, "ausbildungswoche_von"),
        "ausbildungswoche_bis": __read_node(root_node, "ausbildungswoche_bis"),
        "abteilung": __read_node(root_node, "abteilung"),
        "betriebliche_taetigkeit": __read_node(root_node, "betriebliche_taetigkeit"),
        "betrieblicher_unterricht": __read_node(root_node, "betrieblicher_unterricht"),
        "berufsschule": __read_node(root_node, "berufsschule"),
    }
    return input


def __fill_template(input_file_path: Path, output_file_path: Path, template_file_path: Path) -> None:
    shutil.copy(template_file_path, output_file_path)
    document = Document(output_file_path)

    input = __read_input(input_file_path)

    document.tables[0].cell(1, 2).text = f"Ausbildungsnachweis Nr.: {input['nr']}"
    document.tables[0].cell(2, 2).text = f"{input['name']}, {input['vorname']}"
    document.tables[0].cell(0, 5).text = input["ausbildungsjahr"]
    document.tables[0].cell(1, 5).text = input["ausbildungswoche_von"]
    document.tables[0].cell(2, 5).text = input["ausbildungswoche_bis"]
    document.tables[0].cell(3, 5).text = input["abteilung"]
    document.tables[0].cell(5, 5).text = input["betriebliche_taetigkeit"]
    document.tables[0].cell(7, 5).text = input["betrieblicher_unterricht"]
    document.tables[0].cell(9, 5).text = input["berufsschule"]

    document.save(output_file_path)


def __convert_docx_to_pdf(input_file_path: Path, output_dir_path: Path) -> None:
    subprocess.run(
        [LIBREOFFICE_BINARY, "--headless", "--convert-to", "pdf", "--outdir", output_dir_path, input_file_path],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
